#!/usr/bin/env python3
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT


from tqdm.auto import tqdm
import warnings

from chocolatechip.times_config import times_dict
from chocolatechip.MySQLConnector import MySQLConnector
from chocolatechip.intersections import intersection_lookup

warnings.filterwarnings("ignore", message=".*supports SQLAlchemy connectable.*")

# -------------------------
# Word table border utility
# -------------------------
def _set_table_borders_black(table):
    """Force all outer/inner table borders to be black, single-line."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{side}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')  # ~0.5pt
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        borders.append(elem)
    tblPr.append(borders)

# -------------------------
# Time window helpers
# -------------------------
def _windows(iid: int, period: str):
    ts_list = times_dict.get(iid, {}).get(period, [])
    return list(zip(ts_list[0::2], ts_list[1::2]))

# -------------------------
# DB fetch (uses new MySQL method)
# -------------------------
def fetch_conflicts_for_period(db: MySQLConnector, iid: int, p2v: int, period: str) -> pd.DataFrame:
    """
    Concatenate conflicts across all windows in times_dict[iid][period],
    using MySQLConnector.fetchConflictRecordsWithCoordsAndTypeTime.
    """
    wins = _windows(iid, period)
    if not wins:
        return pd.DataFrame(columns=[
            "timestamp","unique_ID1","unique_ID2","cluster1","cluster2",
            "conflict_x","conflict_y","conflict_type","time"
        ])

    label = "P2V" if p2v == 1 else "V2V"
    frames = []
    with tqdm(total=len(wins), desc=f"{iid} {period} {label}", unit="win") as pbar:
        for start, end in wins:
            df = db.fetchConflictRecordsWithCoordsAndTypeTime(
                intersec_id=iid, p2v=p2v, start=start, end=end
            )
            if not df.empty:
                frames.append(df)
                pbar.set_postfix(rows=len(df))
            pbar.update(1)

    if not frames:
        return pd.DataFrame(columns=[
            "timestamp","unique_ID1","unique_ID2","cluster1","cluster2",
            "conflict_x","conflict_y","conflict_type","time"
        ])

    out = (
        pd.concat(frames, ignore_index=True)
        .drop_duplicates()
        .sort_values("timestamp")
        .reset_index(drop=True)
    )
    return out

# -------------------------
# DOCX writing
# -------------------------
def set_font_size(table, size_pt: int = 8):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)


def add_conflict_table(doc: Document, title: str, df: pd.DataFrame):
    """Add a titled table (no orientation changes here)."""
    doc.add_heading(title, level=2)

    if df.empty:
        doc.add_paragraph("No data for this selection.")
        doc.add_paragraph()
        return

    df = df.copy()

    # Pretty timestamp + column order
    if pd.api.types.is_datetime64_any_dtype(df.get("timestamp", pd.Series(dtype="datetime64[ns]"))):
        df["timestamp"] = df["timestamp"].dt.strftime("%Y-%m-%d %H:%M:%S")

    cols = [
        "timestamp",
        "unique_ID1",
        "unique_ID2",
        "cluster1",
        "cluster2",
        "conflict_x",
        "conflict_y",
        "conflict_type",
        "time"
    ]
    cols = [c for c in cols if c in df.columns]
    df = df[cols]

    # Optional: numeric formatting
    for c in ("conflict_x", "conflict_y", "time"):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce").map(lambda v: f"{v:.2f}" if pd.notna(v) else "")

    # Build table
    rows, cols_n = len(df) + 1, len(df.columns)
    table = doc.add_table(rows=rows, cols=cols_n)
    table.style = "Table Grid"
    _set_table_borders_black(table)

    # Center table horizontally
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column width rules
    width_map = {
        "timestamp": Inches(1.5),
        "time": Inches(0.5),
        "conflict_x": Inches(0.75),
        "conflict_y": Inches(0.75),
        "conflict_type": Inches(0.75),
        "unique_ID1": Inches(1.25),
        "unique_ID2": Inches(1.25)
    }

    # Apply widths
    for col_name, width in width_map.items():
        if col_name in df.columns:
            col_idx = df.columns.get_loc(col_name)
            for row in table.rows:
                row.cells[col_idx].width = width

    # Header
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = col

    # Body
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = "" if pd.isna(val) else str(val)

    # Dense font
    set_font_size(table, 8)
    doc.add_paragraph()


def build_doc(out_path: str = "conflict_tables.docx"):
    db = MySQLConnector()
    doc = Document()
    doc.add_heading("Conflict Tables (Landscape, Small Font, with Coordinates & Type/Time)", level=1)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Create a single landscape section for all tables
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    w, h = sec.page_width, sec.page_height
    sec.page_width, sec.page_height = h, w

    intersections = [3248, 3287]
    periods = ["before", "after"]
    kinds = [("P2V", 1), ("V2V", 0)]

    total_tasks = len(intersections) * len(periods) * len(kinds)
    with tqdm(total=total_tasks, desc="All tables", unit="task") as allbar:
        for iid in intersections:
            inter_name = intersection_lookup.get(iid, str(iid))
            doc.add_heading(f"Intersection: {inter_name}", level=2)
            for period in periods:
                for label, p2v in kinds:
                    df = fetch_conflicts_for_period(db, iid=iid, p2v=p2v, period=period)
                    add_conflict_table(doc, title=f"{label} — {period.capitalize()}", df=df)
                    allbar.update(1)

    doc.save(out_path)
    print(f"Written {out_path}")

if __name__ == "__main__":
    build_doc()
