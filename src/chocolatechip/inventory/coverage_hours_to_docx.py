#!/usr/bin/env python3
import pandas as pd
from datetime import datetime, timedelta
from collections import defaultdict, Counter
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from chocolatechip.times_config     import times_dict
from chocolatechip.intersections   import intersection_lookup
from dates_and_daysofweek_counter import abbreviate_name, WEEKDAYS, DAY_ABBR

# Local (not shared) settings for this script
# Hours 7..18 inclusive (displayed as 7 AM–6 PM)
HOURS = list(range(7, 19))
# Threshold: mark an hour as "covered" if at least 20 minutes (1200 sec) recorded
THRESHOLD_SECONDS = 20 * 60

def hour_label(h: int) -> str:
    """Turn 24h int -> 'h AM/PM' safely."""
    suffix = "AM" if h < 12 else "PM"
    hour12 = ((h + 11) % 12) + 1  # 0->12, 13->1, etc.
    return f"{hour12} {suffix}"

def compute_hourly_seconds(ts_list):
    """
    Build per-date, per-hour coverage (in seconds) from a list of [start_iso, end_iso, ...].
    """
    coverage = defaultdict(Counter)  # {date: {hour: seconds}}
    for start_s, end_s in zip(ts_list[0::2], ts_list[1::2]):
        start = datetime.fromisoformat(start_s)
        end   = datetime.fromisoformat(end_s)
        cur   = start
        while cur < end:
            hour_start    = cur.replace(minute=0, second=0, microsecond=0)
            next_boundary = hour_start + timedelta(hours=1)
            seg_end       = min(end, next_boundary)
            coverage[cur.date()][hour_start.hour] += (seg_end - cur).total_seconds()
            cur = seg_end
    return coverage

def build_weekday_hour_table(coverage, threshold: int = THRESHOLD_SECONDS):
    """
    For each weekday and hour in HOURS, count how many *dates* had at least `threshold`
    seconds recorded in that hour.
    """
    table = {wd: {h: 0 for h in HOURS} for wd in WEEKDAYS}
    for date, hour_secs in coverage.items():
        wd = date.strftime("%A")
        for h in HOURS:
            if hour_secs.get(h, 0) >= threshold:
                table[wd][h] += 1
    return table

def make_dataframe(iid, period):
    raw_name = intersection_lookup.get(iid, str(iid))
    name     = abbreviate_name(raw_name)
    ts_list  = times_dict.get(iid, {}).get(period, [])
    if not ts_list:
        return name, None

    coverage = compute_hourly_seconds(ts_list)
    counts   = build_weekday_hour_table(coverage)
    df = pd.DataFrame.from_dict(counts, orient="index", columns=HOURS)
    df.index   = [DAY_ABBR[wd] for wd in df.index]
    df.columns = [hour_label(h) for h in df.columns]
    return name, df

def _set_table_borders_black(table):
    """
    Force all outer/inner table borders to be black, single-line.
    Uses proper <w:tblPr> insertion compatible with python-docx.
    """
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        # insert as first child of <w:tbl>
        tbl.insert(0, tblPr)

    # remove existing tblBorders if present
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{side}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')       # 1/8 pt units (4 ~= 0.5pt)
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        borders.append(elem)
    tblPr.append(borders)

def tables_to_word(iids, periods, out_path="intersection_tables.docx"):
    doc = Document()
    doc.add_heading(
        f"Intersection Coverage Tables (≥20 min per hour counts, {hour_label(HOURS[0])}–{hour_label(HOURS[-1])})",
        level=1
    )

    for iid in iids:
        for period in periods:
            name, df = make_dataframe(iid, period)
            doc.add_heading(f"Intersection {name} — {period.capitalize()}", level=2)
            if df is None:
                doc.add_paragraph("No data for this period.")
                continue

            # Build a Word table: +1 for header row & +1 for weekday column
            rows, cols = df.shape[0] + 1, df.shape[1] + 1
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'  # plain grid
            _set_table_borders_black(table)

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = ""  # top-left empty
            for j, col in enumerate(df.columns, start=1):
                hdr_cells[j].text = col

            # Data rows
            for i, (wd, row) in enumerate(df.iterrows(), start=1):
                row_cells = table.rows[i].cells
                row_cells[0].text = wd
                for j, val in enumerate(row, start=1):
                    row_cells[j].text = str(val)

            doc.add_paragraph()  # space before next table

    doc.save(out_path)
    print(f"Written tables to {out_path}")

if __name__ == "__main__":
    # Example usage: only intersections 3248 & 3287, both before and after
    tables_to_word(iids=[3248, 3287], periods=["before", "after"])
